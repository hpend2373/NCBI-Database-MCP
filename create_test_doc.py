#!/usr/bin/env python3
"""Create a test Word document directly"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create document directory if it doesn't exist
doc_dir = '/Users/minyeop/alphagenome/bio-mcp-blast/WordDocuments'
os.makedirs(doc_dir, exist_ok=True)

# Create a new document
doc = Document()

# Add title
title = doc.add_heading('감상문 - AI와 인간의 미래', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add metadata
doc.core_properties.title = '감상문 - AI와 인간의 미래'
doc.core_properties.author = '홍길동'

# Add date
date_para = doc.add_paragraph('2024년 12월 27일')
date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add introduction
doc.add_heading('서론', level=1)
intro = doc.add_paragraph(
    '최근 AI 기술의 급속한 발전은 우리 사회에 많은 변화를 가져오고 있다. '
    '특히 생성형 AI의 등장은 인간의 창의성과 지능에 대한 근본적인 질문을 던지고 있다.'
)

# Add main body
doc.add_heading('본론', level=1)

doc.add_heading('AI 기술의 현재', level=2)
body1 = doc.add_paragraph(
    'ChatGPT, Claude, Gemini 등의 대규모 언어 모델은 이제 일상적인 도구가 되었다. '
    '이들은 단순한 질문 답변을 넘어 창의적인 글쓰기, 코딩, 분석 등 다양한 분야에서 '
    '인간을 보조하고 있다.'
)

doc.add_heading('인간과 AI의 협업', level=2)
body2 = doc.add_paragraph(
    'AI는 인간을 대체하는 것이 아니라 인간의 능력을 확장시키는 도구로 봐야 한다. '
    '인간의 창의성, 윤리적 판단, 감정적 이해는 여전히 AI가 따라올 수 없는 영역이다.'
)

# Add conclusion
doc.add_heading('결론', level=1)
conclusion = doc.add_paragraph(
    'AI와 인간의 미래는 경쟁이 아닌 협력의 관계로 발전해야 한다. '
    '우리는 AI를 현명하게 활용하면서도 인간 고유의 가치를 지켜나가야 할 것이다.'
)

# Save the document
doc_path = os.path.join(doc_dir, '감상문_AI강연.docx')
doc.save(doc_path)

print(f"✅ 문서 생성 완료: {doc_path}")
print(f"📏 파일 크기: {os.path.getsize(doc_path):,} bytes")

# Also create a simple test document
test_doc = Document()
test_doc.add_heading('Test Document', 0)
test_doc.add_paragraph('This is a test document.')
test_path = os.path.join(doc_dir, 'test.docx')
test_doc.save(test_path)
print(f"✅ 테스트 문서 생성: {test_path}")