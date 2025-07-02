#!/usr/bin/env python3
"""Create a formal Word document with specific formatting requirements"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create document directory if it doesn't exist
doc_dir = '/Users/minyeop/alphagenome/bio-mcp-blast/WordDocuments'
os.makedirs(doc_dir, exist_ok=True)

# Create a new document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

# Add page numbers to footer
from docx.oxml import parse_xml
section = sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Add page number field
run = footer_para.add_run()
fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
instrText = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# Set default style
style = doc.styles['Normal']
font = style.font
font.name = '바탕'
font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '바탕')

# Set paragraph format
paragraph_format = style.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
paragraph_format.line_spacing = 1.15
paragraph_format.space_after = Pt(0)
paragraph_format.space_before = Pt(0)

# Add the content as one continuous paragraph
content = """초불확실성 시대, AI와 나의 역할
KAIST 문술미래전략대학원 서용석 교수님 강연을 들으며 많은 생각을 했다. 평소 미래 기술, 특히 인공지능에 큰 관심을 두고 다양한 기능을 직접 써보며 변화 속도에 감탄하던 나다. 교수님 이야기는 지금 내가 어디에 서 있는지, 앞으로 어디로 가야 할지 중요한 질문을 던졌다. "생성형 AI가 존재 가치를 뿌리부터 흔드는 충격으로 다가왔습니다"라는 교수님의 첫마디는 나 역시 AI 기술을 접하며 느꼈던 경외감과 불안감에 공감을 자아냈다.
교수님은 우리가 초불확실성 시대에 살고 있다고 강조했다. 예상치 못한 사건들을 'X-이벤트'로 정의하며, 기후 변화, 초연결성 확산, 그리고 가장 강력한 요인인 기술 혁신을 그 원인으로 꼽았다. AI 기술 발전 속도는 정말 눈부시다. 교수님이 언급한 오픈AI 딥 리서치 기능은 나에게도 큰 충격이었다. 200달러짜리 기능을 직접 써보진 못했지만, 무료 버전이나 다른 대안들을 통해 비슷한 기능을 접했을 때 이미 충분히 놀랐다. 불과 몇 년 전만 해도 상상하기 어려웠던, 학술 논문 작성을 10분도 안 되는 시간에 처리할 수 있다는 사실은 마치 SF 영화 같았다.
교수님이 딥 리서치를 쓰며 '대학교수라는 직업을 얼마나 더 할 수 있을까?'라고 고민했을 때, 나는 '내 전공은 AI 때문에 어떻게 바뀔까?'라고 비슷하게 생각했다. AI 기술은 이제 단순한 도구를 넘어, 인간 고유 영역이던 창의적이고 복합적인 작업까지 수행한다. 특히 올해는 에이전트의 해라 불릴 만큼 에이전트 AI 발전이 두드러진다. 나도 요즘 클로드 코드 에이전트 같은 기술을 깊이 공부하는데, 코딩은 물론 다양한 자동화 작업을 척척 해내는 것을 보며 감탄한다. 에이전트들이 반복적인 업무를 대신하면서, 우리는 더 고차원적인 사고와 판단에 집중할 수 있을 것이다. 이는 생산성 향상을 넘어 산업 전반에 혁명적인 변화를 가져올 게 분명하다.
이러한 초불확실성 시대에 우리가 갖춰야 할 중요한 역량으로 교수님은 두 가지를 제시했다. 첫째는 회복 탄력성(Resilience)이다. 위기나 재난을 단순히 이전 상태로 복원하는 것을 넘어, 이를 기회 삼아 시스템을 몇 단계 업그레이드시키는 능력이라는 설명이 인상 깊었다. 우리나라가 IMF 외환 위기를 극복하고 금융 시스템을 발전시킨 사례는, 개인 삶에서도 실패와 좌절을 통해 배우고 더 단단해지는 자세가 중요함을 깨닫게 했다. 교수님은 민첩함(Agile)과 회복 탄력성을 합쳐 '에질리언스(Agilience)'라는 개념을 제시했다. 이는 이 시대에 필요한 핵심 역량을 정확히 짚었다고 본다. 나 역시 대학 생활 중 겪었던 실패와 좌절을 통해 성장할 수 있었다. 앞으로도 에질리언스 역량을 꾸준히 길러야겠다고 다짐한다.
두 번째 역량은 미래 문해력이었다. 단순히 글을 읽고 이해하는 것을 넘어, 미래 변화를 예측하고 주도하며 새로운 가능성을 상상하는 능력을 의미한다고 하셨다. 로열 더치 쉘의 시나리오 플래닝과 일론 머스크의 선제적 투자 사례는 미래 문해력이 어떻게 성공으로 이어지는지 보여주었다. 단기 이익보다 장기적 관점에서 미래 사회 흐름을 읽고 대비하는 것이 중요함을 깨달았다. 나는 평소에도 다양한 미래 기술 트렌드를 탐색하며 미래를 예측해보는 것을 즐긴다. 이런 활동들이 내 미래 설계에 중요한 자산이 될 수 있다는 확신을 얻었다. 최근 내년에 AI가 과학적 발견을 할 수도 있다는 기사를 접했는데, 이는 미래 문해력을 기반으로 한 상상이 현실이 될 수 있음을 시사한다. 이러한 변화에 민감하게 반응하고, 새로운 지식을 끊임없이 습득하는 자기 주도적 학습력이 필수적이라는 교수님 말씀이 마음에 와닿았다.
교수님 지도 교수님이 남긴 "Read, learn, engage, humbly act and keep learning!"이라는 문구는 앞으로 내가 걸어가야 할 길에 이정표와 같았다. 끊임없이 배우고, 변화에 능동적으로 대응하며, 실패를 기회로 전환하는 삶의 자세를 통해 초불확실성 시대 속에서도 나만의 가치 있는 역할을 찾아야겠다고 생각한다.
강연 말미, 교수님이 던진 "어쩌면 AI가 우리에게 던진 가장 큰 도전은 바로 나다움이란 무엇인가?라는 질문일지도 모르겠습니다"라는 질문은 큰 울림을 주었다. 기술은 인간 역할을 대체할 수 있어도 인간 그 자체를 대체할 수는 없다는 말씀은 나에게 큰 위로이자 동시에 깊은 성찰의 계기가 됐다. AI가 아무리 뛰어나도 결국 미래를 결정하는 것은 인간이라는 사실을 명심하고, 나만의 고유한 강점과 가치를 발전시켜야 할 때라고 느낀다. 이제 불확실성과 두려움을 넘어 나만의 이야기를 써내려갈 시간이다. 이 강연을 통해 얻은 영감을 바탕으로, 나는 앞으로도 끊임없이 배우고 탐구하며, AI 시대의 변화 속에서 나다운 미래를 만들어갈 것이다."""

# Replace line breaks with spaces to make it one continuous paragraph
content = content.replace('\n', ' ')

# Add the content
para = doc.add_paragraph(content)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Ensure the font is applied to all runs
for run in para.runs:
    run.font.name = '바탕'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '바탕')

# Save the document
doc_path = os.path.join(doc_dir, '감상문1.docx')
doc.save(doc_path)

print(f"✅ 문서 생성 완료: {doc_path}")
print(f"📏 파일 크기: {os.path.getsize(doc_path):,} bytes")
print("\n📋 문서 형식:")
print("- 글자 크기: 10pt")
print("- 글씨체: 바탕체")
print("- 정렬: 양쪽 맞춤")
print("- 줄간격: 1.15")
print("- 여백: 상하 2.54cm, 좌우 1.91cm")
print("- 머리글/바닥글: 0.5cm")
print("- 페이지 번호: 표시됨")
print("- 내용: 한 문단으로 이어서 작성")